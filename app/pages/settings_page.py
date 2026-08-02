#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
单词突围 - 设置页面
==================
"""

import sys, os, threading, io
from datetime import datetime
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

import flet as ft
from app.theme import (
    PRIMARY, PRIMARY_LIGHT, SECONDARY, SURFACE, SUCCESS, ERROR, BACKGROUND,
    TEXT_PRIMARY, TEXT_SECONDARY, TEXT_HINT,
    PAGE_PADDING, CARD_GAP, SPACING_SM, SPACING_MD, SPACING_LG, SPACING_XL, SPACING_XXL,
    RADIUS_MD, RADIUS_SM, RADIUS_XL, RADIUS_LG,
    SHADOW_SM, SHADOW_MD,
    FONT_XS, FONT_SM, FONT_BODY, FONT_LG,
    COLOR_SETTINGS, BORDER, W_MEDIUM, W_SEMIBOLD,
)
from app.components.app_card import AppCard
from app.services import api_service, local_db
from backend.models import StudyRecord, DailyPlan

# 检查Word文档依赖
try:
    from docx import Document
    from docx.shared import Pt
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _clear_study_records():
    s = local_db._get_session()
    try:
        s.query(StudyRecord).delete()
        s.query(DailyPlan).delete()
        s.commit()
        return True
    except:
        s.rollback()
        return False
    finally:
        s.close()


def _generate_docx(words_by_date):
    """生成Word文档 — 含单词/音标/词性/释义/固定搭配/派生词，不含记忆法和例句"""
    doc = Document()
    doc.add_heading('学习记录（导出不含记忆法和例句）', level=0)
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('')

    total = sum(len(words) for words in words_by_date.values())
    for date, words in words_by_date.items():
        doc.add_heading(f'{date}（{len(words)}词）', level=1)
        for w in words:
            p = doc.add_paragraph()
            run = p.add_run(f"{w['word']}")
            run.bold = True
            run.font.size = Pt(12)
            if w.get('phonetic'):
                run2 = p.add_run(f"  {w['phonetic']}")
                run2.font.size = Pt(10)
            if w.get('pos'):
                run3 = p.add_run(f"  [{w['pos']}]")
                run3.font.size = Pt(10)
            doc.add_paragraph(f"释义：{w.get('meaning', '')}", style='List Bullet')
            if w.get('collocations'):
                colls = w['collocations'].replace('|', '；')
                doc.add_paragraph(f"固定搭配：{colls}", style='List Bullet')
            if w.get('derivatives'):
                doc.add_paragraph(f"派生词：{w['derivatives']}", style='List Bullet')
            if w.get('extensions'):
                doc.add_paragraph(f"扩展：{w['extensions']}", style='List Bullet')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), total

def _generate_txt(words_by_date):
    """生成纯文本 — 不含记忆法和例句"""
    lines = []
    lines.append("学习记录（导出不含记忆法和例句）")
    lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append("")

    total = sum(len(words) for words in words_by_date.values())
    for date, words in words_by_date.items():
        lines.append(f"【{date}】（{len(words)}词）")
        lines.append("")
        for w in words:
            lines.append(f"  {w['word']}  {w.get('phonetic', '')}  [{w.get('pos', '')}]")
            lines.append(f"    释义：{w.get('meaning', '')}")
            if w.get('collocations'):
                colls = w['collocations'].replace('|', '；')
                lines.append(f"    固定搭配：{colls}")
            if w.get('derivatives'):
                lines.append(f"    派生词：{w['derivatives']}")
            if w.get('extensions'):
                lines.append(f"    扩展：{w['extensions']}")
            lines.append("")
        lines.append("")
    lines.append(f"总计：{total}词")
    return '\n'.join(lines), total


def _format_date(d_str):
    """2026-07-12 \u2192 7\u670812\u65e5"""
    try:
        dt = datetime.strptime(d_str, '%Y-%m-%d')
        return f"{dt.month}\u6708{dt.day}\u65e5"
    except:
        return d_str


class SettingsPage:
    def __init__(self, app):
        self.app = app
        self.page = app.page
        self._target_value = ft.TextField(
            value=str(api_service.get_daily_target()),
            width=80, height=42,
            text_size=FONT_LG,
            text_align=ft.TextAlign.CENTER,
            keyboard_type=ft.KeyboardType.NUMBER,
            border=ft.InputBorder.OUTLINE,
            border_color=COLOR_SETTINGS,
            focused_border_color=COLOR_SETTINGS,
            content_padding=ft.Padding(left=8, right=8, top=6, bottom=6),
        )
        # 日期选择状态
        self._date_checks = {}    # { date_str: Checkbox }
        self._preview_container = ft.Container(
            content=ft.Text("\u9009\u62e9\u4e0a\u65b9\u65e5\u671f\u67e5\u770b\u5355\u8bcd",
                           size=FONT_SM, color=TEXT_HINT),
            padding=SPACING_MD,
        )
        self._quick_days = ft.TextField(
            value="7", width=60, height=38,
            text_size=14, text_align=ft.TextAlign.CENTER,
            keyboard_type=ft.KeyboardType.NUMBER,
            border=ft.InputBorder.OUTLINE,
            border_color=COLOR_SETTINGS,
            focused_border_color=COLOR_SETTINGS,
            content_padding=ft.Padding(left=8, right=8, top=4, bottom=4),
        )
        self._export_format = "docx"
        self._download_btn = ft.Container(
            content=ft.Text("\u4e0b\u8f7d\u6587\u6863", color=ft.Colors.WHITE,
                           size=14, weight=ft.FontWeight.BOLD),
            padding=ft.Padding(24, 12, 24, 12),
            bgcolor=ft.Colors.GREY_400,
            border_radius=RADIUS_XL,
            ink=True,
            on_click=None,
        )

    def build(self):
        # 加载学习日期
        dates = []
        try:
            dates = api_service.get_study_dates()
        except:
            dates = []
        self._date_checks = {}

        date_rows = []
        for d_idx, d in enumerate(dates):
            cb = ft.Checkbox(
                label=f"{_format_date(d['date'])}\uff08{d['count']}\u8bcd\uff09",
                value=False,
                on_change=self._on_date_toggle,
                active_color=COLOR_SETTINGS,
                check_color=ft.Colors.WHITE,
            )
            self._date_checks[d['date']] = cb
            date_rows.append(
                ft.Container(
                    content=cb,
                    padding=ft.Padding(left=8, top=6, right=8, bottom=6),
                    border_radius=RADIUS_SM,
                    bgcolor=ft.Colors.with_opacity(0.03, COLOR_SETTINGS) if d_idx % 2 == 1 else None,
                )
            )

        if not date_rows:
            date_rows = [
                ft.Container(
                    content=ft.Text("\u6682\u65e0\u5b66\u4e60\u8bb0\u5f55",
                                   size=FONT_SM, color=TEXT_HINT),
                    padding=SPACING_MD,
                )
            ]

        # 预览区域（初始空）
        self._preview_container.content = ft.Column([
            ft.Text("\u9009\u62e9\u65e5\u671f\u540e\u9884\u89c8\u5355\u8bcd",
                   size=FONT_SM, color=TEXT_HINT),
        ], spacing=0)

        # 关于信息
        about_items = [
            ("\u5e94\u7528", "\u5355\u8bcd\u7a81\u56f4 (\u72ec\u7acb\u7248)"),
            ("\u7248\u672c", "2.1.0"),
            ("\u6570\u636e\u6765\u6e90", "\u300a\u5355\u8bcd\u7a81\u56f45200\u300b\u4e0a\u518c"),
            ("\u590d\u4e60\u7b97\u6cd5", "\u827e\u5bbe\u6d69\u65af\u9057\u5fd8\u66f2\u7ebf"),
            ("\u5355\u8bcd\u603b\u6570", "2281 \u4e2a"),
            ("\u8fd0\u884c\u6a21\u5f0f", "\u672c\u5730\u6570\u636e\u5e93 (\u65e0\u9700\u540e\u7aef)"),
        ]
        about_rows = []
        for i, (label, value) in enumerate(about_items):
            about_rows.append(self._info_row(label, value, i))
            if i < len(about_items) - 1:
                about_rows.append(ft.Divider(height=1, color=BORDER))

        return ft.Container(
            content=ft.ListView([
                # ================================================================
                # 每日学习目标
                # ================================================================
                AppCard(
                    content=ft.Column([
                        ft.Row([
                            ft.Container(
                                content=ft.Icon(ft.Icons.TRACK_CHANGES,
                                            color=COLOR_SETTINGS, size=18),
                            padding=ft.Padding(6, 6, 6, 6),
                            bgcolor=ft.Colors.with_opacity(0.10, COLOR_SETTINGS),
                            border_radius=8,
                        ),
                        ft.Container(width=8),
                        ft.Text("\u6bcf\u65e5\u5b66\u4e60\u76ee\u6807", size=FONT_LG,
                                weight=ft.FontWeight.BOLD, color=TEXT_PRIMARY),
                    ]),
                    ft.Divider(height=1, color=BORDER),
                    ft.Container(height=SPACING_MD),
                    ft.Row([
                        ft.Text("\u6bcf\u5929\u65b0\u5b66\u5355\u8bcd", size=FONT_BODY,
                                color=TEXT_SECONDARY, weight=W_MEDIUM),
                        ft.Container(expand=True),
                        self._target_value,
                    ]),
                    ft.Container(height=SPACING_MD),
                    ft.Row([
                        ft.Container(expand=True),
                        ft.Container(
                            content=ft.Row([
                                ft.Icon(ft.Icons.SAVE, color=ft.Colors.WHITE, size=16),
                                ft.Container(width=4),
                                ft.Text("\u4fdd\u5b58\u76ee\u6807", color=ft.Colors.WHITE,
                                        size=14, weight=ft.FontWeight.BOLD),
                            ]),
                            padding=ft.Padding(20, 10, 20, 10),
                            gradient=ft.LinearGradient(
                                colors=["#64748B", "#94A3B8"],
                                begin=ft.Alignment(-1, 0),
                                end=ft.Alignment(1, 0),
                            ),
                            border_radius=RADIUS_MD,
                            ink=True,
                            on_click=self._save_target,
                        ),
                    ]),
                    ft.Container(height=4),
                    ft.Text("\u5efa\u8bae\uff1a\u6bcf\u592910-20\u4e2a\uff0c\u6709\u57fa\u7840\u53ef30-50\u4e2a",
                            size=FONT_SM, color=TEXT_HINT, italic=True),
                ], spacing=0),
                elevation="sm",
            ),

            ft.Container(height=CARD_GAP),

            # ================================================================
            # 学习记录
            # ================================================================
            AppCard(
                content=ft.Column([
                    ft.Row([
                        ft.Container(
                            content=ft.Icon(ft.Icons.DATE_RANGE,
                                            color=COLOR_SETTINGS, size=18),
                            padding=ft.Padding(6, 6, 6, 6),
                            bgcolor=ft.Colors.with_opacity(0.10, COLOR_SETTINGS),
                            border_radius=8,
                        ),
                        ft.Container(width=8),
                        ft.Text("\u5b66\u4e60\u8bb0\u5f55", size=FONT_LG,
                                weight=ft.FontWeight.BOLD, color=TEXT_PRIMARY),
                    ]),
                    ft.Divider(height=1, color=BORDER),
                    ft.Container(height=SPACING_SM),
                    # 快捷导出
                    ft.Container(
                        content=ft.Column([
                            ft.Text("快捷导出", size=FONT_SM,
                                    color=TEXT_SECONDARY, weight=W_MEDIUM),
                            ft.Container(height=6),
                            ft.Row([
                                ft.Container(
                                    content=ft.Text("近3天", color=ft.Colors.WHITE,
                                                   size=12, weight=ft.FontWeight.BOLD),
                                    padding=ft.Padding(16, 8, 16, 8),
                                    bgcolor=COLOR_SETTINGS,
                                    border_radius=RADIUS_XL,
                                    ink=True,
                                    on_click=lambda e, d=3: self._quick_export(d),
                                ),
                                ft.Container(
                                    content=ft.Text("近5天", color=ft.Colors.WHITE,
                                                   size=12, weight=ft.FontWeight.BOLD),
                                    padding=ft.Padding(16, 8, 16, 8),
                                    bgcolor=COLOR_SETTINGS,
                                    border_radius=RADIUS_XL,
                                    ink=True,
                                    on_click=lambda e, d=5: self._quick_export(d),
                                ),
                                ft.Container(
                                    content=ft.Text("近7天", color=ft.Colors.WHITE,
                                                   size=12, weight=ft.FontWeight.BOLD),
                                    padding=ft.Padding(16, 8, 16, 8),
                                    bgcolor=COLOR_SETTINGS,
                                    border_radius=RADIUS_XL,
                                    ink=True,
                                    on_click=lambda e, d=7: self._quick_export(d),
                                ),
                                ft.Container(
                                    content=ft.Text("自定义", color=COLOR_SETTINGS,
                                                   size=12, weight=ft.FontWeight.BOLD),
                                    padding=ft.Padding(12, 8, 12, 8),
                                    bgcolor=ft.Colors.with_opacity(0.10, COLOR_SETTINGS),
                                    border_radius=RADIUS_XL,
                                    ink=True,
                                    on_click=self._quick_export_custom,
                                ),
                                self._quick_days,
                                ft.Text("天", size=FONT_SM, color=TEXT_HINT),
                            ], spacing=8, wrap=True),
                        ], spacing=0),
                    ),
                    ft.Container(height=SPACING_SM),
                    # 日期列表（可多选）
                    *([ft.Container(
                        content=ft.Column([
                            ft.Text("\u9009\u62e9\u65e5\u671f", size=FONT_SM,
                                    color=TEXT_SECONDARY, weight=W_MEDIUM),
                            ft.Container(height=4),
                            *date_rows,
                        ], spacing=0),
                    )] if date_rows else [
                        ft.Container(
                            content=ft.Text("\u6682\u65e0\u5b66\u4e60\u8bb0\u5f55",
                                           size=FONT_SM, color=TEXT_HINT),
                            padding=SPACING_MD,
                        )
                    ]),
                    ft.Divider(height=1, color=BORDER),
                    ft.Container(height=SPACING_SM),
                    ft.Divider(height=1, color=BORDER),
                    ft.Container(height=SPACING_SM),
                    # 导出格式切换
                    ft.Container(
                        content=ft.Row([
                            ft.Text("导出格式", size=FONT_SM,
                                    color=TEXT_SECONDARY, weight=W_MEDIUM),
                            ft.Container(width=12),
                            ft.Container(
                                content=ft.Text("Word", color=COLOR_SETTINGS,
                                               size=12, weight=ft.FontWeight.BOLD),
                                padding=ft.Padding(12, 6, 12, 6),
                                border=ft.Border.all(1.5, COLOR_SETTINGS),
                                border_radius=RADIUS_MD,
                                ink=True,
                                on_click=lambda e: self._set_format("docx"),
                            ),
                            ft.Container(width=6),
                            ft.Container(
                                content=ft.Text("纯文本", color=TEXT_HINT,
                                               size=12, weight=ft.FontWeight.BOLD),
                                padding=ft.Padding(12, 6, 12, 6),
                                border=ft.Border.all(1.5, COLOR_SETTINGS),
                                border_radius=RADIUS_MD,
                                ink=True,
                                on_click=lambda e: self._set_format("txt"),
                            ),
                        ]),
                    ),
                    ft.Container(height=SPACING_SM),
                    # 预览区域
                    ft.Container(
                        content=ft.Column([
                            ft.Text("\u5355\u8bcd\u9884\u89c8", size=FONT_SM,
                                    color=TEXT_SECONDARY, weight=W_MEDIUM),
                            ft.Container(height=4),
                            ft.Container(
                                content=self._preview_container,
                                border=ft.Border(
                                    left=ft.BorderSide(3, COLOR_SETTINGS),
                                ),
                                border_radius=RADIUS_SM,
                                padding=ft.Padding(
                                    left=SPACING_SM, top=0, right=0, bottom=0
                                ),
                                bgcolor=ft.Colors.with_opacity(0.02, COLOR_SETTINGS),
                            ),
                        ], spacing=0),
                    ),
                    ft.Container(height=SPACING_MD),
                    # 下载按钮
                    ft.Container(
                        content=self._download_btn,
                    ),
                ], spacing=0),
                elevation="sm",
            ),

            ft.Container(height=CARD_GAP),

            # ================================================================
            # 数据管理
            # ================================================================
            AppCard(
                content=ft.Column([
                    ft.Row([
                        ft.Container(
                            content=ft.Icon(ft.Icons.STORAGE, color="#FF8F00", size=18),
                            padding=ft.Padding(6, 6, 6, 6),
                            bgcolor=ft.Colors.with_opacity(0.10, "#FF8F00"),
                            border_radius=8,
                        ),
                        ft.Container(width=8),
                        ft.Text("\u6570\u636e\u7ba1\u7406", size=FONT_LG,
                                weight=ft.FontWeight.BOLD, color=TEXT_PRIMARY),
                    ]),
                    ft.Divider(height=1, color=BORDER),
                    ft.Container(height=SPACING_MD),
                    ft.Container(
                        content=ft.Row([
                            ft.Container(
                                content=ft.Icon(ft.Icons.DELETE_SWEEP,
                                                color=ERROR, size=20),
                                padding=ft.Padding(4, 4, 4, 4),
                                bgcolor=ft.Colors.with_opacity(0.10, ERROR),
                                border_radius=6,
                            ),
                            ft.Container(width=10),
                            ft.Column([
                                ft.Text("\u91cd\u7f6e\u5b66\u4e60\u8bb0\u5f55",
                                        size=FONT_BODY, color=ERROR, weight=W_SEMIBOLD),
                                ft.Text("\u6e05\u7a7a\u6240\u6709\u5b66\u4e60\u8bb0\u5f55\uff0c"
                                        "\u5355\u8bcd\u6570\u636e\u4e0d\u53d7\u5f71\u54cd",
                                        size=FONT_XS, color=TEXT_HINT),
                            ], spacing=2, expand=True),
                            ft.Container(
                                content=ft.Icon(ft.Icons.CHEVRON_RIGHT,
                                                color=ERROR, size=20),
                            ),
                        ]),
                        padding=ft.Padding(12, 10, 12, 10),
                        bgcolor=ft.Colors.with_opacity(0.06, ERROR),
                        border=ft.Border.all(1, ft.Colors.with_opacity(0.15, ERROR)),
                        border_radius=RADIUS_SM,
                        ink=True,
                        on_click=self._confirm_reset,
                    ),
                ], spacing=0),
                elevation="sm",
            ),

            ft.Container(height=CARD_GAP),

            # ================================================================
            # 启动帮助
            # ================================================================
            AppCard(
                content=ft.Column([
                    ft.Row([
                        ft.Container(
                            content=ft.Icon(ft.Icons.PLAY_CIRCLE_OUTLINE,
                                            color=SECONDARY, size=18),
                            padding=ft.Padding(6, 6, 6, 6),
                            bgcolor=ft.Colors.with_opacity(0.10, SECONDARY),
                            border_radius=8,
                        ),
                        ft.Container(width=8),
                        ft.Text("\u4e0b\u6b21\u5982\u4f55\u6253\u5f00", size=FONT_LG,
                                weight=ft.FontWeight.BOLD, color=TEXT_PRIMARY),
                    ]),
                    ft.Divider(height=1, color=BORDER),
                    ft.Container(height=8),
                    ft.Text("\u53cc\u51fb start.bat \u6216\u5728\u7ec8\u7aef\u8fd0\u884c\uff1a",
                            size=FONT_SM, color=TEXT_SECONDARY),
                    ft.Container(height=4),
                    ft.Container(
                        content=ft.Text(
                            "cd E:\\APP\npython run_app.py",
                            size=13, color="#34D399", font_family="monospace",
                            weight=W_MEDIUM,
                        ),
                        padding=ft.Padding(left=16, top=12, right=16, bottom=12),
                        bgcolor="#1E293B",
                        border_radius=RADIUS_SM,
                    ),
                    ft.Container(height=8),
                    ft.Text("\u6d4f\u89c8\u5668\u6253\u5f00 http://localhost:8551",
                            size=FONT_SM, color=TEXT_SECONDARY),
                ], spacing=0),
                elevation="sm",
            ),

            ft.Container(height=CARD_GAP),

            # ================================================================
            # 关于
            # ================================================================
            AppCard(
                content=ft.Column([
                    ft.Row([
                        ft.Container(
                            content=ft.Icon(ft.Icons.INFO, color=COLOR_SETTINGS, size=18),
                            padding=ft.Padding(6, 6, 6, 6),
                            bgcolor=ft.Colors.with_opacity(0.10, COLOR_SETTINGS),
                            border_radius=8,
                        ),
                        ft.Container(width=8),
                        ft.Text("\u5173\u4e8e", size=FONT_LG,
                                weight=ft.FontWeight.BOLD, color=TEXT_PRIMARY),
                    ]),
                    ft.Divider(height=1, color=BORDER),
                    ft.Container(height=8),
                    *about_rows,
                ], spacing=0),
                elevation="sm",
            ),

            ft.Container(height=SPACING_LG),
            ], spacing=0),
            padding=ft.Padding(left=PAGE_PADDING, top=PAGE_PADDING,
                               right=PAGE_PADDING, bottom=0),
        )

    def _info_row(self, label, value, index=0):
        is_alt = index % 2 == 1
        return ft.Container(
            content=ft.Row([
                ft.Text(label, size=FONT_BODY, color=TEXT_SECONDARY, weight=W_MEDIUM),
                ft.Container(expand=True),
                ft.Text(value, size=FONT_BODY, color=TEXT_PRIMARY, weight=W_SEMIBOLD),
            ]),
            padding=ft.Padding(left=SPACING_MD, top=10, right=SPACING_MD, bottom=10),
            bgcolor=ft.Colors.with_opacity(0.04, COLOR_SETTINGS) if is_alt else None,
        )

    def _save_target(self, e):
        try:
            val = int(self._target_value.value)
            if val < 1:
                self.app.show_snackbar("\u6bcf\u65e5\u76ee\u6807\u81f3\u5c11\u4e3a1", ERROR)
                return
            if val > 200:
                self.app.show_snackbar("\u76ee\u6807\u592a\u9ad8\u4e86\uff0c\u5efa\u8bae\u4e0d\u8d85\u8fc750", ERROR)
                return
            ok = api_service.set_daily_target(val)
            if ok:
                self.app.show_snackbar(f"\u6bcf\u65e5\u65b0\u8bcd\u76ee\u6807\u5df2\u8bbe\u4e3a {val} \u4e2a")
            else:
                self.app.show_snackbar("\u4fdd\u5b58\u5931\u8d25", ERROR)
        except ValueError:
            self.app.show_snackbar("\u8bf7\u8f93\u5165\u6709\u6548\u6570\u5b57", ERROR)

    def _on_date_toggle(self, e):
        """日期选择变化 → 刷新预览和下载按钮"""
        selected = [d for d, cb in self._date_checks.items() if cb.value]
        if not selected:
            self._preview_container.content = ft.Column([
                ft.Text("\u9009\u62e9\u65e5\u671f\u540e\u9884\u89c8\u5355\u8bcd",
                       size=FONT_SM, color=TEXT_HINT),
            ], spacing=0)
            self._download_btn.on_click = None
            self._download_btn.bgcolor = ft.Colors.GREY_400
            self.page.update()
            return

        # 异步加载预览
        threading.Thread(target=self._load_preview, args=(selected,), daemon=True).start()

    def _on_date_toggle_has_selection(self):
        return any(cb.value for cb in self._date_checks.values())

    def _load_preview(self, selected):
        """加载选中日期的单词预览"""
        try:
            words_by_date = api_service.get_words_by_dates(selected)
        except Exception:
            words_by_date = {}

        total = sum(len(v) for v in words_by_date.values())
        preview_rows = []
        for date in sorted(words_by_date.keys()):
            words = words_by_date[date]
            preview_rows.append(
                ft.Container(
                    content=ft.Text(
                        f"{_format_date(date)}\uff08{len(words)}\u8bcd\uff09",
                        size=FONT_SM, weight=ft.FontWeight.BOLD,
                        color=TEXT_PRIMARY,
                    ),
                    padding=ft.Padding(top=4, bottom=2),
                )
            )
            for w in words[:5]:  # 最多显示5个
                preview_rows.append(
                    ft.Container(
                        content=ft.Row([
                            ft.Container(
                                width=6, height=6,
                                bgcolor=COLOR_SETTINGS,
                                border_radius=3,
                            ),
                            ft.Container(width=6),
                            ft.Text(w['word'], size=FONT_SM, color=TEXT_PRIMARY,
                                    expand=True),
                            ft.Text(
                                "\u2713" if w.get('result') == 'remember' else "\u2717",
                                size=FONT_XS,
                                color=SUCCESS if w.get('result') == 'remember' else ERROR,
                            ),
                        ], spacing=0),
                        padding=ft.Padding(left=2, top=2, right=2, bottom=2),
                    )
                )
            if len(words) > 5:
                preview_rows.append(
                    ft.Text(
                        f"... \u8fd8\u6709{len(words)-5}\u8bcd",
                        size=FONT_XS, color=TEXT_HINT,
                    )
                )

        if not preview_rows:
            preview_rows = [ft.Text("\u6682\u65e0\u6570\u636e", size=FONT_SM, color=TEXT_HINT)]

        self._preview_container.content = ft.Column(
            preview_rows, spacing=2, scroll=ft.ScrollMode.AUTO
        )
        # 启用下载按钮
        self._download_btn.on_click = lambda e: self._do_download(selected)
        self._download_btn.bgcolor = COLOR_SETTINGS
        self._download_btn.update()
        self._preview_container.update()

    def _do_download(self, selected, words_by_date=None):
        """按勾选日期导出（服务端路由下载，手机端可靠）"""
        if not selected:
            self.app.show_snackbar("请先勾选要导出的日期", ERROR)
            return
        fmt = "txt" if self._export_format == "txt" else "docx"
        self._trigger_export(fmt, dates=selected)

    def _trigger_export(self, fmt, days=None, dates=None):
        """构建 /export 下载链接并弹出对话框（TextSpan url 导航，不走 data URL）"""
        if fmt == "docx" and not _DOCX_AVAILABLE:
            self.app.show_snackbar(
                "缺少 python-docx 库，运行 pip install python-docx", ERROR)
            return
        params = [f'fmt={fmt}']
        if days:
            params.append(f'days={days}')
        elif dates:
            params.append('dates=' + ','.join(dates))
        url = '/export?' + '&'.join(params)

        link = ft.Text(
            spans=[ft.TextSpan(
                text="⬇️ 点击这里下载文件",
                url=url,
                style=ft.TextStyle(size=16, color=PRIMARY,
                                   weight=ft.FontWeight.BOLD),
            )],
        )
        dlg = ft.AlertDialog(
            title=ft.Text("导出学习记录"),
            content=ft.Column([
                link,
                ft.Container(height=6),
                ft.Text("如果未自动下载，点上方链接即可保存文件",
                        size=12, color=TEXT_HINT),
            ], tight=True, spacing=4),
            actions=[ft.TextButton(
                "关闭", on_click=lambda e: self.app.close_dialog(dlg))],
            shape=ft.RoundedRectangleBorder(radius=RADIUS_MD),
        )
        self.page.overlay.append(dlg)
        dlg.open = True
        self.page.update()

    def _quick_export(self, days):
        """快捷导出最近N天（服务端路由下载）"""
        fmt = "txt" if self._export_format == "txt" else "docx"
        self._trigger_export(fmt, days=days)
        self.app.show_snackbar(f"已生成近{days}天导出文件")

    def _quick_export_custom(self, e):
        """自定义天数导出"""
        try:
            days = int(self._quick_days.value)
            if days < 1:
                self.app.show_snackbar("天数至少为1", ERROR)
                return
            if days > 365:
                self.app.show_snackbar("最多365天", ERROR)
                return
            self._quick_export(days)
        except ValueError:
            self.app.show_snackbar("请输入有效数字", ERROR)

    def _set_format(self, fmt):
        """切换导出格式"""
        self._export_format = fmt
        label = "Word" if fmt == "docx" else "纯文本"
        text = f"下载 {label}" if fmt == "docx" else "下载 纯文本"
        self._download_btn.content = ft.Container(
            content=ft.Text(text, color=ft.Colors.WHITE,
                           size=14, weight=ft.FontWeight.BOLD),
            padding=ft.Padding(24, 12, 24, 12),
            bgcolor=COLOR_SETTINGS if self._on_date_toggle_has_selection() else ft.Colors.GREY_400,
            border_radius=RADIUS_XL,
            ink=True,
            on_click=self._download_btn.on_click,
        )
        self._download_btn.update()

    def _confirm_reset(self, e):
        dlg = ft.AlertDialog(
            title=ft.Text("\u786e\u8ba4\u91cd\u7f6e"),
            content=ft.Text("\u786e\u5b9a\u8981\u6e05\u9664\u6240\u6709\u5b66\u4e60\u8bb0\u5f55\u5417\uff1f"
                            "\u5355\u8bcd\u6570\u636e\u4e0d\u4f1a\u4e22\u5931\u3002"),
            actions=[
                ft.TextButton("\u53d6\u6d88",
                    on_click=lambda e: self.app.close_dialog(dlg)),
                ft.TextButton("\u786e\u5b9a\u91cd\u7f6e",
                    style=ft.ButtonStyle(color=ERROR),
                    on_click=lambda e: self._do_reset(dlg)),
            ],
        )
        self.page.overlay.append(dlg)
        dlg.open = True
        self.page.update()

    def _do_reset(self, dlg):
        self.app.close_dialog(dlg)
        ok = _clear_study_records()
        if ok:
            self.app.show_snackbar("\u5b66\u4e60\u8bb0\u5f55\u5df2\u6e05\u7a7a")
        else:
            self.app.show_snackbar("\u91cd\u7f6e\u5931\u8d25")
